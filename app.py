"""
DocuMind Backend - Production Ready Version
Fixed ChromaDB compatibility for Python 3.13
"""

import os
import sys
import uuid
import json
import tempfile
import traceback
import threading
from datetime import datetime
from typing import List, Dict, Optional
from dataclasses import dataclass, asdict

from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
import PyPDF2
import docx

# Load .env file
env_path = os.path.join(os.path.dirname(__file__), ".env")
if os.path.exists(env_path):
    with open(env_path, "r") as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                key, value = line.split("=", 1)
                os.environ[key.strip()] = value.strip().strip('"').strip("'")

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
if not GROQ_API_KEY:
    print("WARNING: GROQ_API_KEY not set. Create a .env file with GROQ_API_KEY=gsk_your_key")
    GROQ_API_KEY = "missing"
else:
    print("API key loaded. Length:", len(GROQ_API_KEY))

CHROMA_DB_PATH = os.path.join(os.path.dirname(__file__), "chroma_db")
DOCUMENTS_METADATA_FILE = os.path.join(os.path.dirname(__file__), "documents_metadata.json")
COLLECTION_NAME = "documind_documents"
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200

@dataclass
class Document:
    id: str
    name: str
    type: str
    size: str
    content: str
    chunks: List[str]
    uploaded_at: str

@dataclass
class ChatMessage:
    role: str
    content: str
    sources: Optional[List[Dict]] = None
    timestamp: str = None
    def __post_init__(self):
        if self.timestamp is None: self.timestamp = datetime.now().isoformat()

@dataclass
class Conversation:
    id: str
    title: str
    messages: List[ChatMessage]
    document_ids: List[str]
    created_at: str
    def __post_init__(self):
        if isinstance(self.created_at, str) and not self.created_at:
            self.created_at = datetime.now().isoformat()

class DocumentProcessor:
    @staticmethod
    def extract_text(file_path: str, file_type: str) -> str:
        if file_type.lower() == "pdf": return DocumentProcessor._extract_pdf(file_path)
        elif file_type.lower() in ["docx", "doc"]: return DocumentProcessor._extract_docx(file_path)
        elif file_type.lower() == "txt": return DocumentProcessor._extract_txt(file_path)
        else: raise ValueError("Unsupported: " + file_type)
    @staticmethod
    def _extract_pdf(file_path: str) -> str:
        text = ""
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages):
                pt = page.extract_text()
                if pt: text += "\n--- Page " + str(i+1) + " ---\n" + pt
        return text
    @staticmethod
    def _extract_docx(file_path: str) -> str:
        return "\n".join([p.text for p in docx.Document(file_path).paragraphs])
    @staticmethod
    def _extract_txt(file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f: return f.read()
    @staticmethod
    def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[str]:
        chunks, start = [], 0
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            if end < len(text):
                bp = max(chunk.rfind("."), chunk.rfind("\n"))
                if bp > chunk_size * 0.5:
                    end = start + bp + 1
                    chunk = text[start:end]
            chunks.append(chunk.strip())
            start = end - overlap
        return chunks

class VectorStore:
    def __init__(self):
        self._client = None
        self._collection = None
        self._embedding_function = None
        self._initialized = False
        self._error = None

    def _init(self):
        if self._initialized: return True
        if self._error: return False
        try:
            print("Initializing ChromaDB at:", CHROMA_DB_PATH)
            import chromadb
            from chromadb.utils import embedding_functions

            # Ensure directory exists
            os.makedirs(CHROMA_DB_PATH, exist_ok=True)

            # Use Client with explicit settings to avoid tenant issues
            self._client = chromadb.PersistentClient(
                path=CHROMA_DB_PATH,
                settings=chromadb.Settings(
                    anonymized_telemetry=False,
                    allow_reset=True
                )
            )
            print("   ChromaDB client created")

            print("   Loading embedding model (may take 30-60s on first run)...")
            # Load embedding model with timeout protection
            model_loaded = False
            model_error = None
            
            def load_model():
                nonlocal model_loaded, model_error
                try:
                    self._embedding_function = embedding_functions.SentenceTransformerEmbeddingFunction(
                        model_name="all-MiniLM-L6-v2"
                    )
                    model_loaded = True
                except Exception as e:
                    model_error = e
            
            # Load model in a thread with timeout
            thread = threading.Thread(target=load_model, daemon=True)
            thread.start()
            thread.join(timeout=120)  # 2 minute timeout
            
            if not model_loaded:
                if model_error:
                    raise RuntimeError(f"Model loading failed: {model_error}")
                else:
                    raise RuntimeError("Model loading timed out after 120 seconds. Check your internet connection and try restarting the app.")
            
            print("   Model loaded")

            # Get or create collection
            try:
                self._collection = self._client.get_collection(
                    name=COLLECTION_NAME,
                    embedding_function=self._embedding_function
                )
                print("   Existing collection found")
            except Exception:
                self._collection = self._client.create_collection(
                    name=COLLECTION_NAME,
                    embedding_function=self._embedding_function,
                    metadata={"hnsw:space": "cosine"}
                )
                print("   New collection created")

            self._initialized = True
            return True
        except Exception as e:
            self._error = str(e)
            print("   ChromaDB Error:", e)
            traceback.print_exc()
            return False

    def add_document(self, doc_id: str, chunks: List[str], metadata: Dict):
        if not self._init():
            raise RuntimeError("Vector store failed: " + str(self._error))
        chunk_ids = [doc_id + "_chunk_" + str(i) for i in range(len(chunks))]
        metas = [{"doc_id": doc_id, "doc_name": metadata.get("name", "Unknown"), "chunk_index": i, "total_chunks": len(chunks), **metadata} for i, _ in enumerate(chunks)]
        self._collection.add(ids=chunk_ids, documents=chunks, metadatas=metas)

    def search(self, query: str, n_results: int = 5) -> List[Dict]:
        if not self._init(): return []
        results = self._collection.query(query_texts=[query], n_results=n_results, include=["documents", "metadatas", "distances"])
        out = []
        if results.get("ids") and len(results["ids"][0]) > 0:
            for i in range(len(results["ids"][0])):
                out.append({"chunk_id": results["ids"][0][i], "text": results["documents"][0][i], "metadata": results["metadatas"][0][i], "distance": results["distances"][0][i]})
        return out

    def delete_document(self, doc_id: str):
        if self._init():
            try: self._collection.delete(where={"doc_id": doc_id})
            except Exception as e: print("Delete warning:", e)

class AIEngine:
    def __init__(self, api_key: str):
        if not api_key or api_key == "missing": raise ValueError("API key required")
        from groq import Groq
        self.client = Groq(api_key=api_key)
        self.model = "llama-3.3-70b-versatile"
    def generate_response(self, query: str, context: List[Dict], conversation_history=None) -> Dict:
        ctx, sources = "", []
        for i, chunk in enumerate(context):
            ctx += "\n\n[Doc " + str(i+1) + ": " + chunk["metadata"].get("doc_name", "?") + "]\n" + chunk["text"]
            sn = chunk["metadata"].get("doc_name", "?")
            if sn not in [s["name"] for s in sources]: sources.append({"name": sn, "relevance": round((1-chunk["distance"])*100, 1)})
        prompt = "You are DocuMind. Answer STRICTLY from the context below. If not found, say so. Cite sources.\n\nCONTEXT:\n" + ctx + "\n\nQUESTION: " + query
        msgs = [{"role": "system", "content": prompt}]
        if conversation_history:
            for msg in conversation_history[-6:]: msgs.append({"role": msg["role"], "content": msg["content"]})
        msgs.append({"role": "user", "content": query})
        try:
            r = self.client.chat.completions.create(model=self.model, messages=msgs, temperature=0.3, max_tokens=2048)
            return {"text": r.choices[0].message.content, "sources": sources[:3]}
        except Exception as e: return {"text": "Error: " + str(e), "sources": []}
    def generate_summary(self, text: str) -> str:
        try:
            r = self.client.chat.completions.create(model=self.model, messages=[{"role": "system", "content": "Summarize in 3-5 points."}, {"role": "user", "content": text[:4000]}], temperature=0.3, max_tokens=1024)
            return r.choices[0].message.content
        except Exception as e: return "Error: " + str(e)

class DocuMindApp:
    def __init__(self):
        self.documents = {}
        self.conversations = {}
        self.processor = DocumentProcessor()
        self.vector_store = VectorStore()
        self.ai_engine = None
        
        # Load existing documents from disk
        self._load_documents()
        
        # Eagerly initialize vector store at startup
        print("Initializing Vector Store...")
        if not self.vector_store._init():
            print("   WARNING: Vector store initialization failed. Uploads may not work.")
        else:
            print("   Vector Store ready")
            # Clean up any orphaned documents from before persistence was added
            self._cleanup_orphaned_documents()
        
        if GROQ_API_KEY and GROQ_API_KEY != "missing":
            try:
                print("Initializing AI engine...")
                self.ai_engine = AIEngine(GROQ_API_KEY)
                print("   AI ready")
            except Exception as e:
                print("   AI failed:", e)
        else:
            print("   No API key. AI disabled.")
    
    def _load_documents(self):
        """Load documents metadata from disk"""
        if os.path.exists(DOCUMENTS_METADATA_FILE):
            try:
                with open(DOCUMENTS_METADATA_FILE, 'r') as f:
                    docs_data = json.load(f)
                    for doc_id, doc_data in docs_data.items():
                        doc = Document(
                            id=doc_data['id'],
                            name=doc_data['name'],
                            type=doc_data['type'],
                            size=doc_data['size'],
                            content=doc_data['content'],
                            chunks=doc_data['chunks'],
                            uploaded_at=doc_data['uploaded_at']
                        )
                        self.documents[doc_id] = doc
                    if self.documents:
                        print(f"   Loaded {len(self.documents)} existing documents from disk")
            except Exception as e:
                print(f"   WARNING: Could not load documents: {e}")
    
    def _cleanup_orphaned_documents(self):
        """Remove documents from ChromaDB that aren't in our metadata (from before persistence was added)"""
        try:
            if not self.vector_store._init():
                return
            
            # Get all document IDs from ChromaDB
            all_data = self.vector_store._collection.get()
            if not all_data['ids']:
                return
            
            orphaned_doc_ids = set()
            for chunk_id in all_data['ids']:
                # Extract doc_id from chunk_id (format: doc_id_chunk_0)
                doc_id = chunk_id.rsplit('_chunk_', 1)[0]
                if doc_id not in self.documents:
                    orphaned_doc_ids.add(doc_id)
            
            if orphaned_doc_ids:
                print(f"   Cleaning up {len(orphaned_doc_ids)} orphaned document(s) from ChromaDB...")
                for doc_id in orphaned_doc_ids:
                    self.vector_store.delete_document(doc_id)
                print(f"   Cleanup complete")
        except Exception as e:
            print(f"   WARNING: Could not cleanup orphaned documents: {e}")
    
    def _save_documents(self):
        """Save documents metadata to disk"""
        try:
            docs_data = {}
            for doc_id, doc in self.documents.items():
                docs_data[doc_id] = asdict(doc)
            with open(DOCUMENTS_METADATA_FILE, 'w') as f:
                json.dump(docs_data, f, indent=2)
        except Exception as e:
            print(f"   WARNING: Could not save documents: {e}")
    def upload_document(self, file: UploadFile):
        doc_id = str(uuid.uuid4())
        ext = file.filename.split(".")[-1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix="." + ext) as tmp:
            content = file.file.read(); tmp.write(content); tmp_path = tmp.name
        try:
            raw = self.processor.extract_text(tmp_path, ext)
            chunks = self.processor.chunk_text(raw)
            size = str(round(len(content)/1024, 1)) + " KB" if len(content) < 1024*1024 else str(round(len(content)/(1024*1024), 1)) + " MB"
            doc = Document(id=doc_id, name=file.filename, type=ext.upper(), size=size, content=raw[:5000], chunks=chunks, uploaded_at=datetime.now().isoformat())
            self.vector_store.add_document(doc_id=doc_id, chunks=chunks, metadata={"name": file.filename, "type": ext, "uploaded_at": doc.uploaded_at})
            self.documents[doc_id] = doc
            self._save_documents()  # Persist to disk
            return doc
        finally: os.unlink(tmp_path)
    def chat(self, cid: str, msg: str):
        if not self.ai_engine: return {"response": "AI not initialized. Set GROQ_API_KEY in .env and restart.", "sources": [], "conversation_id": cid}
        if cid not in self.conversations: self.conversations[cid] = Conversation(id=cid, title="New", messages=[], document_ids=[], created_at=datetime.now().isoformat())
        conv = self.conversations[cid]
        conv.messages.append(ChatMessage(role="user", content=msg))
        results = self.vector_store.search(msg, n_results=5)
        hist = [{"role": m.role, "content": m.content} for m in conv.messages[:-1]]
        resp = self.ai_engine.generate_response(msg, results, hist)
        conv.messages.append(ChatMessage(role="assistant", content=resp["text"], sources=resp["sources"]))
        return {"response": resp["text"], "sources": resp["sources"], "conversation_id": cid}
    def get_summary(self, doc_id: str):
        if not self.ai_engine: return "AI not available"
        if doc_id not in self.documents: raise ValueError("Not found")
        return self.ai_engine.generate_summary(" ".join(self.documents[doc_id].chunks))
    def delete_document(self, doc_id: str):
        if doc_id in self.documents: del self.documents[doc_id]
        self.vector_store.delete_document(doc_id)
        self._save_documents()  # Persist to disk
    def list_documents(self):
        return [{"id": d.id, "name": d.name, "type": d.type, "size": d.size, "uploaded_at": d.uploaded_at} for d in self.documents.values()]

app = FastAPI(title="DocuMind", version="1.0.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])

print("\n" + "="*60)
print("Starting DocuMind...")
print("="*60)
try:
    app_instance = DocuMindApp()
    print("\nReady!")
except Exception as e:
    print("\nInit failed:", e)
    traceback.print_exc()
    app_instance = None
print("="*60 + "\n")

@app.post("/api/upload")
async def upload(file: UploadFile = File(...)):
    if not app_instance: raise HTTPException(503, "Initializing")
    try:
        d = app_instance.upload_document(file)
        return JSONResponse({"success": True, "document": {"id": d.id, "name": d.name, "type": d.type, "size": d.size, "uploaded_at": d.uploaded_at}})
    except Exception as e: raise HTTPException(500, str(e))

@app.post("/api/chat")
async def chat(message: str = Form(...), conversation_id: str = Form(default="default")):
    if not app_instance: raise HTTPException(503, "Initializing")
    try:
        r = app_instance.chat(conversation_id, message)
        return JSONResponse({"success": True, "response": r["response"], "sources": r["sources"], "conversation_id": r["conversation_id"]})
    except Exception as e: raise HTTPException(500, str(e))

@app.get("/api/documents")
async def list_docs():
    if not app_instance: return JSONResponse({"success": True, "documents": []})
    return JSONResponse({"success": True, "documents": app_instance.list_documents()})

@app.delete("/api/documents/{doc_id}")
async def delete(doc_id: str):
    if not app_instance: raise HTTPException(503, "Initializing")
    try: app_instance.delete_document(doc_id); return JSONResponse({"success": True})
    except Exception as e: raise HTTPException(500, str(e))

@app.get("/api/documents/{doc_id}/summary")
async def summary(doc_id: str):
    if not app_instance: raise HTTPException(503, "Initializing")
    try: return JSONResponse({"success": True, "summary": app_instance.get_summary(doc_id)})
    except Exception as e: raise HTTPException(500, str(e))

@app.get("/api/health")
async def health():
    return JSONResponse({"status": "healthy" if app_instance else "init", "timestamp": datetime.now().isoformat(), "docs": len(app_instance.documents) if app_instance else 0, "ai_ready": bool(app_instance.ai_engine) if app_instance else False})

if __name__ == "__main__":
    import socket
    import uvicorn

    def find_free_port(start: int = 8000, end: int = 8010) -> int:
        for port in range(start, end + 1):
            with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
                s.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
                try:
                    s.bind(("0.0.0.0", port))
                    return port
                except OSError:
                    continue
        raise RuntimeError(f"No free port found between {start} and {end}")

    port = int(os.getenv('DOCUMIND_PORT', '8000'))
    chosen_port = port
    try:
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as s:
            s.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
            s.bind(("0.0.0.0", port))
    except OSError:
        chosen_port = find_free_port(start=8000, end=8010)
        print(f"Port {port} is already in use. Falling back to {chosen_port}.")

    print("="*60)
    print(f"DocuMind Server - http://localhost:{chosen_port}")
    print("="*60)
    uvicorn.run(app, host="0.0.0.0", port=chosen_port)
